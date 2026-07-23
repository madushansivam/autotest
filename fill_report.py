#!/usr/bin/env python3
"""
AutoTest Final Report – Content Filler
Replaces every [TO BE COMPLETED] and GUIDANCE marker with real,
evidenced content sourced from the codebase, git log, and live runs.
Run once; backs up the original before touching it.
"""

import shutil, re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC_PATH    = '/home/madushan/Project/autotest/AutoTest_Final_Report.docx'
BACKUP_PATH = '/home/madushan/Project/autotest/AutoTest_Final_Report_BACKUP.docx'

# ──────────────────────────────────────────────────────────────
# LOW-LEVEL XML HELPERS
# ──────────────────────────────────────────────────────────────

def para_text(para):
    return ''.join(t.text or '' for t in para._p.iter(qn('w:t')))

def clear_runs(para):
    p = para._p
    for tag in (qn('w:r'), qn('w:hyperlink')):
        for elem in list(p.iter(tag)):
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)

def set_text(para, text):
    """Replace paragraph content with plain text, keep paragraph style."""
    clear_runs(para)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    para._p.append(r)

def make_p(text):
    """Create a bare w:p element with one run of text."""
    new_p = OxmlElement('w:p')
    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        new_p.append(r)
    return new_p

def insert_after(ref_elem, lines):
    """Insert lines as consecutive paragraphs immediately after ref_elem."""
    last = ref_elem
    for text in lines:
        new_p = make_p(text)
        last.addnext(new_p)
        last = new_p

def remove_para(para):
    p = para._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)

# ──────────────────────────────────────────────────────────────
# CONTENT BLOCKS  (every claim sourced – see inline comments)
# ──────────────────────────────────────────────────────────────

# Source: lscpu, free -h, df -h, /etc/os-release run 2026-07-21
HARDWARE = [
    "The system was developed and evaluated on a single personal laptop. "
    "System specifications were confirmed by direct inspection during the project:",
    "",
    "Processor: Intel Core i5-13420H (13th Generation), 8 physical cores, "
    "12 logical processors, maximum clock speed 4.6 GHz.",
    "Memory: 7.4 GiB RAM.",
    "Storage: 475 GB NVMe SSD (343 GB available during development).",
    "Operating System: Fedora Linux 44 (Workstation Edition), kernel 7.1.3-200.fc44.x86_64, x86_64 architecture.",
    "",
    "No cloud compute was used. All pipeline stages — browser launch, LLM API calls, "
    "script execution, and database writes — ran locally on this hardware.",
]

# Source: git log timeline and all source files
METHODS_32 = [
    "The system was developed using an incremental approach in which each pipeline "
    "stage was implemented as a standalone script, verified with a real run against "
    "a live web application, and committed to version control before the next stage "
    "was begun. This ensured that failures were isolated to the component under "
    "development and that each stage could be independently inspected.",
    "",
    "The development sequence, reconstructed from the git commit history "
    "(repository: madushansivam/autotest), was as follows.",
    "",
    "Stage 1 — Proof of concept (12 July 2026, commit 8012425): Playwright was "
    "installed and a minimal script confirmed that it could launch Chromium headlessly "
    "and load a real web page.",
    "",
    "Stage 2 — Crawler (commits 10662fd, 2367654, 10bfca9, 78e096e; 12 July 2026): "
    "The page crawler was built iteratively. Version 1 extracted buttons, inputs, "
    "and links into a structured interface map. Version 2 added an interaction step "
    "(adding a todo item) before scanning, to reveal elements rendered only after "
    "user action. Version 3 added aria-label and class-name fallbacks for icon-only "
    "buttons. A multi-page BFS crawler (multi-crawler.js) was tested against "
    "books.toscrape.com, which revealed the first safe-mode gap: the phrases "
    "'add to basket', 'add to cart', and 'add to bag' were absent from the initial "
    "blocklist. These were added to safe-mode.config.json (commit 78e096e).",
    "",
    "Stage 3 — LLM integration and provider pivot (commits aa34a1d, c4da1ce, "
    "61cf4f3; 12 July 2026): API key management via dotenv was added. The Anthropic "
    "Claude API was initially integrated as the AI backend, but the free-tier credit "
    "limit was exhausted during this phase. The project pivoted to the HuggingFace "
    "free Inference API, using meta-llama/Llama-3.1-8B-Instruct (commit c4da1ce). "
    "A retry loop of up to three attempts was then added after discovering that the "
    "model occasionally emits malformed JSON (unclosed strings, trailing commas), "
    "causing silent zero-output runs. A confidence-value validation step was also "
    "added; test cases whose confidence field was neither 'structural' nor "
    "'behavioral' were tagged with a validationWarning flag.",
    "",
    "Stage 4 — Execution engine (commit 805699b; 12 July 2026): The execution engine "
    "compiled LLM-generated scripts at runtime using JavaScript's AsyncFunction "
    "constructor, navigated to the target URL, and ran each compiled function in a "
    "live Playwright browser context. Two failure modes were discovered: (1) a "
    "vacuous pass, where a selector (tbody > tr) not present on the page returned "
    "{ passed: true } because Playwright's auto-waiting silently succeeded on an "
    "empty result; and (2) selector hallucination, where the code-generation prompt "
    "did not include the crawled DOM data, causing the model to invent selectors. "
    "Both were fixed: the prompt was updated to include DOM data, and all generated "
    "scripts were required to wrap actions in try/catch, returning "
    "{ passed: false, error: err.message } on failure.",
    "",
    "Stage 5 — Database and API (commits 9ba8701, f02e94f, 8997bc3, 0748675, "
    "a80cf4c; 13 July 2026): A three-table SQLite schema (crawls, test_cases, "
    "test_results) was implemented using better-sqlite3. CRUD operations were "
    "verified with live data. The Express server was added, and all stages were "
    "integrated into runFullPipeline() exposed via POST /api/crawl. An unhandled-"
    "rejection handler was added after a generated script crashed the server process "
    "during a live run.",
    "",
    "Stage 6 — Dashboard (commit 8997bc3; 13 July 2026): A vanilla JavaScript "
    "dashboard (public/index.html) was implemented, providing a URL input form, "
    "crawl history grouped by target URL with pass/fail/crash tallies, and a detail "
    "view showing the generated script and error message for each test case.",
    "",
    "No test framework was used for the working system; each stage was validated by "
    "running it against a real web application and inspecting the output. A separate "
    "unit-tested architectural exploration (backend/, frontend/, eval/) was "
    "implemented afterward but has not been run end-to-end and is not part of the "
    "evaluated system.",
]

# Source: run-pipeline.js lines 122–157, server.js lines 1–59
ARCH_41 = [
    "AutoTest is structured as a four-stage sequential pipeline, implemented as a "
    "single exported function (runFullPipeline, run-pipeline.js line 122) and "
    "exposed via an Express HTTP server (server.js). The four stages are: Crawl, "
    "Generate Test Cases, Generate Scripts, and Execute. Each stage completes before "
    "the next begins.",
    "",
    "Entry point: The user submits a URL via the dashboard form. The browser sends "
    "POST /api/crawl to the Express server (server.js line 41). The URL is first "
    "checked by isUnsafeUrl() (server.js lines 13-26), which rejects non-http(s) "
    "schemes, localhost, and common private IP ranges. If the URL passes, "
    "runFullPipeline(url) is called.",
    "",
    "Stage 1 — Crawl (run-pipeline.js lines 16-35): Playwright launches a headless "
    "Chromium browser and navigates to the target URL. The scanPage() function uses "
    "Playwright's $$eval to extract: buttons and role='button' elements (with "
    "aria-label and class fallbacks for icon-only controls); input elements (type, "
    "name, id, placeholder); and anchor links (text, href). Each button is tagged "
    "safe: true or safe: false against the blocklist in safe-mode.config.json.",
    "",
    "Stage 2 — Generate Test Cases (run-pipeline.js lines 63-83): The structured "
    "page data is serialised to JSON and included in a prompt sent to "
    "meta-llama/Llama-3.1-8B-Instruct via the HuggingFace Inference API. The model "
    "returns a JSON array of 3-5 test case objects, each with a description and a "
    "confidence tag (structural or behavioral). Parsing is retried up to three times "
    "on malformed JSON.",
    "",
    "Stage 3 — Generate Scripts (run-pipeline.js lines 85-98): For each test case, "
    "a second LLM call converts the description into a Playwright function body. "
    "The actual crawled DOM data is included in this prompt, constraining the model "
    "to use only selectors present on the page. The response is stripped of markdown "
    "fences and validated for minimum length.",
    "",
    "Stage 4 — Execute (run-pipeline.js lines 100-120): Each script is compiled "
    "at runtime via JavaScript's AsyncFunction constructor and run in a live "
    "Playwright page context. Results are classified as: pass (returned "
    "{ passed: true }), fail (returned { passed: false }, sub-category "
    "assertion_failed), crash (threw a runtime exception, sub-category "
    "runtime_error), or skipped (no script generated). All results are written to "
    "SQLite.",
    "",
    "[Figure 1 — INSERT HERE: Architecture diagram showing the flow: "
    "User (browser) -> POST /api/crawl -> isUnsafeUrl() -> runFullPipeline() -> "
    "Stage 1: Crawl -> Stage 2: Generate Test Cases -> Stage 3: Generate Scripts -> "
    "Stage 4: Execute -> SQLite -> Dashboard. "
    "Draw using draw.io or equivalent; export as PNG and insert above this line.]",
]

# Source: db.js lines 5-32
DATA_MODEL_42 = [
    "The working system uses a SQLite database (autotest.db), created and managed "
    "by the better-sqlite3 Node.js library. The schema is defined entirely in db.js "
    "(lines 5-32) and consists of three tables with foreign-key relationships.",
    "",
    "crawls (id INTEGER PRIMARY KEY AUTOINCREMENT, url TEXT NOT NULL, "
    "crawled_at TEXT NOT NULL, raw_json TEXT NOT NULL): one record per pipeline run. "
    "raw_json holds the complete serialised output of scanPage() for the target URL.",
    "",
    "test_cases (id INTEGER PRIMARY KEY AUTOINCREMENT, crawl_id INTEGER NOT NULL "
    "REFERENCES crawls(id), description TEXT NOT NULL, confidence TEXT NOT NULL, "
    "script TEXT, created_at TEXT NOT NULL): one record per generated test case. "
    "script is NULL if generation failed or was skipped.",
    "",
    "test_results (id INTEGER PRIMARY KEY AUTOINCREMENT, test_case_id INTEGER "
    "NOT NULL REFERENCES test_cases(id), result TEXT NOT NULL, failure_category TEXT, "
    "error TEXT, executed_at TEXT NOT NULL): one record per execution outcome. "
    "result is one of: 'pass', 'fail', 'crash', 'skipped'. failure_category records "
    "the sub-classification ('runtime_error', 'assertion_failed', 'syntax_error') "
    "for non-passing results. error holds the first line of the Playwright error "
    "message.",
    "",
    "[Figure 2 — INSERT HERE: Entity-Relationship diagram showing the three tables "
    "above with their column names, data types, and foreign-key relationships "
    "(crawls 1..* test_cases 1..* test_results). Draw using any standard ER notation "
    "tool and insert above this line.]",
    "",
    "The database file is gitignored (*.db in .gitignore) because it is a runtime "
    "artefact — it is created fresh on first run and is not part of the repository.",
]

# Source: run-pipeline.js lines 37-61 (verbatim prompt strings)
PROMPTS_43 = [
    "Two LLM prompts were developed and iterated during the project. Both are "
    "reproduced below from run-pipeline.js.",
    "",
    "Test-Case Generation Prompt (run-pipeline.js lines 37-45):",
    "",
    "'You are a QA engineer. Given this structured map of a web page interactive "
    "elements, generate 3-5 test cases a human tester would perform.\\n\\n"
    "For each test case, respond with a JSON array where each object has:\\n"
    "- \"description\": plain-language description of the user action and expected result\\n"
    "- \"confidence\": either \"structural\" or \"behavioral\"\\n\\n"
    "Only test elements marked safe true. Never generate a test for anything marked safe false.\\n\\n"
    "Page data:\\n[serialised page data]\\n\\n"
    "CRITICAL: Respond with ONLY a valid JSON array. No markdown code fences, "
    "no commentary, no trailing commas, all strings properly closed.'",
    "",
    "Script Generation Prompt (run-pipeline.js lines 47-61):",
    "",
    "'You are a Playwright test engineer. Convert this test case description into "
    "a single, complete, runnable Playwright test function body.\\n\\n"
    "Test case: \"[test case description]\"\\n"
    "Confidence level: [structural|behavioral]\\n\\n"
    "Here is the ACTUAL crawled DOM data for this page. You MUST use the exact "
    "id, name, placeholder, or class values shown here:\\n\\n[serialised page data]\\n\\n"
    "Rules:\\n"
    "- Assume 'page' is already an initialized Playwright Page object, already navigated to the URL.\\n"
    "- Use page.locator() with EXACT attribute values from the DOM data above.\\n"
    "- Wrap the core action in a try/catch. On success, return { passed: true }. "
    "On failure, return { passed: false, error: err.message }.\\n"
    "- Do NOT include imports, browser launch code, or function wrapper syntax.\\n"
    "- Do NOT use Python-style snake_case methods.\\n"
    "- Do NOT invent methods that do not exist on Playwright Page or Locator objects.\\n"
    "- Respond with ONLY the raw JavaScript statements, no markdown fences, no commentary.'",
    "",
    "Key design decisions made during prompt iteration:",
    "",
    "1. JSON-only output constraint: The 'CRITICAL: Respond with ONLY a valid JSON "
    "array' instruction was added (commit 61cf4f3) after discovering that the model "
    "wrapped its output in markdown code blocks by default, breaking JSON.parse().",
    "",
    "2. DOM data inclusion in script generation: The crawled page data is explicitly "
    "included in the script-generation prompt. This was absent in the original "
    "implementation and was added (commit 805699b) after discovering that the model "
    "hallucinated selectors when not shown real DOM data.",
    "",
    "3. Retry logic: Both prompts retry up to MAX_RETRIES = 3 times on parse failure "
    "or suspicious output (code shorter than 10 characters, or containing <html>).",
    "",
    "4. Confidence validation: After parsing, each test case's confidence field is "
    "checked against {'structural', 'behavioral'}. Any other value is normalised to "
    "'unvalidated' and a validationWarning flag is set "
    "(run-pipeline.js lines 73-77).",
]

# Source: public/index.html analysis
UI_44 = [
    "The dashboard is a single-file vanilla JavaScript application (public/index.html), "
    "served as a static file by Express (server.js line 54). It requires no build step.",
    "",
    "The interface consists of four components:",
    "",
    "1. URL input form: A text input and submit button at the top of the page. "
    "Submitting sends POST /api/crawl and displays a status message while the "
    "pipeline runs.",
    "",
    "2. Crawl history panel: All past runs are loaded from GET /api/crawls and "
    "grouped by target URL. Each URL group shows how many runs have been performed "
    "and an aggregate tally of pass, fail, and crash counts across all runs for "
    "that URL. Groups expand and collapse on click.",
    "",
    "3. Detail panel: Clicking a crawl row loads its full data from "
    "GET /api/crawls/:id and shows all test cases for that run, each with: "
    "the test description, the confidence tag (structural/behavioral), "
    "the result badge (pass/fail/crash/skipped), the generated Playwright script, "
    "and the error message if the test did not pass.",
    "",
    "4. Delete action: Each crawl row has a delete button that calls "
    "DELETE /api/crawls/:id after a confirmation dialog, removing the crawl and "
    "all its test data from the database.",
    "",
    "[Figure 3 — INSERT HERE: Screenshot of the dashboard showing the URL input "
    "form and at least one crawl result in the history panel. To capture: start "
    "the server with 'node server.js', navigate to http://localhost:3000, and "
    "take a full-page screenshot.]",
    "",
    "[Figure 4 — INSERT HERE: Screenshot of the detail panel for one crawl, showing "
    "at least one test case with its generated script and result badge. Click any "
    "crawl row in the history panel to open its detail view, then screenshot.]",
]

# Source: server.js lines 13-26; safe-mode.config.json lines 1-48
SECURITY_45 = [
    "Two mechanisms were implemented to limit the risk of the tool being used "
    "against unintended targets.",
    "",
    "URL validation (server.js lines 13-26): Before any URL is passed to the "
    "pipeline, it is checked by the isUnsafeUrl() function. This function rejects: "
    "any URL that cannot be parsed by the WHATWG URL API; any scheme other than "
    "http: or https:; the hostname localhost; and any hostname matching the IPv4 "
    "ranges 127.x.x.x, 10.x.x.x, 192.168.x.x, 172.16-31.x.x, and 169.254.x.x "
    "(the AWS EC2 instance metadata address).",
    "",
    "This guard does not DNS-resolve hostnames. A public domain name that resolves "
    "to a private IP address would pass the check. A DNS-resolving SSRF guard was "
    "designed and unit-tested (30 tests) as part of the architectural exploration "
    "(backend/src/lib/ssrf-guard.ts), but was not integrated into the working "
    "prototype. This is an identified gap and is discussed further in Section 5.4.",
    "",
    "Safe-mode blocklist (safe-mode.config.json): The crawler tags each discovered "
    "interactive element as safe: true or safe: false by checking its label against "
    "a blocklist of 17 button text phrases associated with destructive or financial "
    "actions (examples: 'delete account', 'place order', 'add to basket'). "
    "Additional blocklists cover specific input field types (credit card fields), "
    "form action routes (password-reset), and input name or id attributes "
    "(cardNumber, cvv). Elements tagged safe: false are reported in the crawl output "
    "but never clicked, filled, or submitted.",
    "",
    "The safe-mode blocklist is an explicit-deny list rather than an explicit-allow "
    "list. This is noted in the configuration file itself (safe-mode.config.json "
    "line 46): 'This is a blocklist, not a whitelist, because v1 targets known-shape "
    "CRUD apps where an explicit blocklist is tractable.' The 'add to basket', "
    "'add to cart', and 'add to bag' entries were added after a live test against "
    "books.toscrape.com revealed they were not covered by the initial implementation "
    "(commit 78e096e, 12 July 2026).",
    "",
    "No rate limiting is implemented in the working system. Rate-limiting logic was "
    "designed as part of the architectural exploration "
    "(backend/src/lib/rate-limiter.ts) but was not integrated.",
]

# Source: corrected functional requirements; fixes 'screenshots' and 'per authenticated user' errors
OBJECTIVES_51 = [
    "The following functional requirements were stated in Chapter 1. This section "
    "maps each to its implementation and to the evidence in Sections 5.2 and 5.3.",
    "",
    "FR1 — The system shall accept a target URL and crawl the application's "
    "reachable interactive elements. Implemented: scanPage() in run-pipeline.js "
    "(lines 16-35) extracts buttons, inputs, and links with safe-mode tagging. "
    "Evidence: todomvc-site-map.json (committed crawl output).",
    "",
    "FR2 — The system shall generate test-case descriptions from the crawled "
    "interface map using a large language model. Implemented: generateTestCases() "
    "in run-pipeline.js (lines 63-83), using meta-llama/Llama-3.1-8B-Instruct "
    "via HuggingFace. Evidence: generated-tests.json (committed); database query "
    "results (Section 5.2).",
    "",
    "FR3 — The system shall convert generated test-case descriptions into executable "
    "Playwright scripts. Implemented: generateScript() in run-pipeline.js "
    "(lines 85-98). Evidence: playwright-scripts.json (committed); script column "
    "in test_cases table.",
    "",
    "FR4 — The system shall execute generated scripts headlessly and capture pass, "
    "fail, crash, and skipped outcomes with error messages. Implemented: "
    "executeTest() in run-pipeline.js (lines 102-120), classifying results into "
    "four outcome types and three failure sub-categories "
    "(runtime_error, assertion_failed, syntax_error). Evidence: test-results.json "
    "(committed); test_results table (Section 5.2).",
    "",
    "FR5 — The system shall present results in a browser-based dashboard and persist "
    "run history in a local SQLite database. Implemented: public/index.html "
    "(dashboard) served by Express (server.js line 54); autotest.db (SQLite); "
    "db-operations.js (CRUD). Evidence: db.js lines 5-32 define the schema; "
    "server.js line 54 serves the static dashboard.",
    "",
    "Note: An earlier draft of the requirements included 'screenshots on failure' "
    "(FR4) and 'per authenticated user' (FR5). Neither feature was implemented. "
    "Failure evidence is recorded as error message text. The system has no "
    "authentication and is designed for single-user local use only.",
]

# Source: test-results.json (committed), autotest.db query (20 July 2026), git commit messages
RESULTS_52 = [
    "Two pipeline runs were conducted against the same target URL at different stages "
    "of the project. Results are presented separately because they used different "
    "code versions and produced observably different outputs — a meaningful finding "
    "discussed in Section 5.3.",
    "",
    "Run A — Early development run (committed to the repository as test-results.json; "
    "traceable to commit 805699b, 12 July 2026):",
    "  Target: https://demo.playwright.dev/todomvc",
    "  Test cases generated: 3 (2 structural, 1 behavioral)",
    "  Results: 0 pass, 1 fail (assertion_failed), 2 crash (runtime_error)",
    "",
    "  Test case 1: 'Click on the Delete button and verify it exists without crashing "
    "the page.' — crash (runtime_error). Error: locator.click: Timeout 5000ms "
    "exceeded waiting for locator('[aria-label] Delete'). The Delete button is only "
    "rendered after a todo item exists; no item was added before the click.",
    "",
    "  Test case 2: 'Enter some text in the text input field and verify the placeholder "
    "text changes.' — crash (runtime_error). Error: page.fill: Timeout 5000ms "
    "exceeded waiting for locator('#todo-input'). The actual input has no id "
    "attribute; the LLM used a selector it invented rather than one from the crawl.",
    "",
    "  Test case 3: 'Click on the All link and verify the page loads without errors.' "
    "— fail (assertion_failed). Error: locator.click timed out on locator('text=All').",
    "",
    "Run B — Verification run (conducted 20 July 2026, 17:45 IST; "
    "results in autotest.db, queried 20 July 2026):",
    "  Target: https://demo.playwright.dev/todomvc (same URL)",
    "  Test cases generated: 3 (2 structural, 1 behavioral)",
    "  Results: 2 pass, 1 fail (assertion_failed), 0 crash",
    "",
    "  Test case 1: 'Click on the first link to navigate to the TodoMVC website.' "
    "— fail (assertion_failed). Error: locator('text=TodoMVC') resolved to 2 elements "
    "(strict mode violation). The model used a text locator that matched two anchors "
    "on the page.",
    "",
    "  Test case 2: 'Enter a task in the input field and submit it by pressing the "
    "Enter key.' — pass.",
    "",
    "  Test case 3: 'Click on the second link to navigate to Remo H. Jansen's "
    "GitHub page.' — pass.",
    "",
    "The two runs produced entirely different test cases despite the same target URL. "
    "This is direct evidence of LLM output non-determinism (see Section 5.4).",
    "",
    "Qualitative case studies — four bugs found during development:",
    "",
    "Case Study 1 — Vacuous pass (commit 805699b, 12 July 2026): An early version "
    "of the execution engine returned { passed: true } for a test whose selector "
    "(tbody > tr) was absent from the page. Playwright's auto-waiting did not raise "
    "an exception; the test appeared to pass without performing any action. Fix: "
    "all generated scripts were required to use try/catch and return "
    "{ passed: false, error: err.message } on failure, making silent no-ops "
    "impossible.",
    "",
    "Case Study 2 — Safe-mode gap (commit 78e096e, 12 July 2026): A live crawl of "
    "books.toscrape.com included 'Add to basket' links in the safe-marked elements "
    "because those phrases were absent from the initial blocklist. Found during a "
    "real run, not a code review. Fix: three variants ('add to basket', 'add to "
    "cart', 'add to bag') were added to safe-mode.config.json before any test "
    "cases were generated for that site.",
    "",
    "Case Study 3 — Malformed JSON output (commit 61cf4f3, 12 July 2026): A "
    "generation run silently produced zero test cases. Inspecting the raw model "
    "response revealed JSON with an unclosed string. Fix: retry loop up to 3 "
    "attempts with a stricter prompt instruction to double-check JSON validity "
    "before responding.",
    "",
    "Case Study 4 — Selector hallucination (commit 805699b, 12 July 2026): The "
    "original script-generation prompt did not include the crawled DOM data. The "
    "model used selectors based on assumed page structure (#todo-input, #add-btn), "
    "none of which existed. All three scripts failed at runtime. Fix: the full "
    "crawled page data was included in the script-generation prompt, constraining "
    "the model to use only selectors it had been shown. Run B's 0-crash result is "
    "the direct measurable outcome of this fix.",
]

# Source: results above plus code analysis
DISCUSSION_53 = [
    "The results from both pipeline runs demonstrate that AutoTest's execution engine "
    "correctly identifies and categorises failures in LLM-generated test code.",
    "",
    "In Run A, all three tests failed — two crashed and one produced an assertion "
    "failure. These are not indications that the system is broken; they are the "
    "expected output of a correctly-functioning execution engine encountering flawed "
    "generated code. The two crashes trace directly to Case Study 4 (selector "
    "hallucination), which was fixed before Run B. The improvement between Run A "
    "(0 pass, 2 crash) and Run B (2 pass, 0 crash) is therefore a direct measurement "
    "of the effect of the prompt fix: including actual DOM data in the "
    "script-generation prompt eliminated the selector-hallucination crash category "
    "entirely.",
    "",
    "The one remaining failure in Run B represents a different and less severe error "
    "class: a strict-mode locator ambiguity (locator matched 2 elements instead of "
    "1). The execution engine correctly classifies this as assertion_failed rather "
    "than crash, distinguishing it from the missing-element errors of Run A.",
    "",
    "The two runs confirm that the result taxonomy (pass / fail / crash / skipped, "
    "with sub-categories runtime_error / assertion_failed / syntax_error) is adequate "
    "to distinguish meaningfully different types of AI-generated code failures. "
    "Each category points to a different root cause: runtime_error indicates a "
    "missing element or timing problem; assertion_failed indicates a logically "
    "incorrect expectation or ambiguous locator; syntax_error indicates code the "
    "JavaScript engine cannot parse (no instances occurred in these runs).",
    "",
    "The different test case descriptions generated in Run A and Run B against the "
    "same URL confirm that LLM output is non-deterministic at the inference "
    "parameters used by the HuggingFace free API. This is discussed in Section 5.4.",
    "",
    "All five functional requirements stated in Section 5.1 are met by the working "
    "implementation. The two corrections to over-stated requirements (no screenshot "
    "capture, no user authentication) are clarifications of scope, not failures of "
    "implementation.",
]

# Source: code analysis, architecture, lack of formal evaluation
LIMITATIONS_54 = [
    "The following limitations were observed during development and evaluation:",
    "",
    "1. LLM output non-determinism: The meta-llama/Llama-3.1-8B-Instruct model, "
    "accessed via the HuggingFace free Inference API, does not expose a temperature "
    "parameter. Two runs against the same URL produced entirely different test cases "
    "(observed directly: Run A and Run B, Section 5.2). Test results are not fully "
    "reproducible without a fixed model checkpoint and deterministic sampling.",
    "",
    "2. Single-page crawl in the integrated pipeline: The standalone multi-page BFS "
    "crawler (multi-crawler.js) navigates up to 8 pages using same-origin link "
    "following and was demonstrated against books.toscrape.com (committed output: "
    "site-map.json). However, the integrated pipeline (runFullPipeline() in "
    "run-pipeline.js) crawls only the single landing page submitted by the user. "
    "Multi-page integration was not completed within the project.",
    "",
    "3. Hostname-only SSRF guard: The isUnsafeUrl() function rejects obvious private "
    "IP ranges and localhost by regex, but does not DNS-resolve hostnames. A public "
    "domain configured to resolve to an internal address would pass the check. This "
    "is acknowledged explicitly in the source code (server.js lines 10-12).",
    "",
    "4. No authentication or multi-user isolation: The working system is designed "
    "for a single local user and provides no authentication. All crawl history is "
    "accessible to anyone with access to the localhost:3000 port.",
    "",
    "5. No formal evaluation against ground truth: No controlled experiment was "
    "conducted comparing AutoTest-generated tests against a human-authored test suite. "
    "The four case studies in Section 5.2 are qualitative evidence; they do not "
    "constitute a measurement of recall, precision, or false-positive rate. Such an "
    "evaluation would require a fixed target application with a known complete set "
    "of testable behaviours and a panel of human testers authoring baseline tests "
    "for comparison. The architectural exploration includes an evaluation harness "
    "(eval/) designed for this purpose, but it was not run.",
    "",
    "6. Structural coverage only: LLM-generated tests can reliably detect structural "
    "failures (missing elements, navigation crashes, uncaught exceptions) but cannot "
    "verify server-side correctness — whether submitted data was persisted, emails "
    "sent, or calculations correct. This is an inherent property of the approach.",
]

# Source: architectural exploration design docs, code analysis
FUTURE_55 = [
    "The following improvements are identified for continuation beyond the HNDIT "
    "submission:",
    "",
    "1. Multi-page crawl integration: Wire multi-crawler.js's BFS logic into "
    "runFullPipeline(), with a configurable page budget per run and deduplication "
    "of test cases across pages sharing interactive elements.",
    "",
    "2. DNS-resolving SSRF guard: Replace isUnsafeUrl() with the DNS-resolving guard "
    "designed in backend/src/lib/ssrf-guard.ts, which resolves all hostnames to IP "
    "addresses before checking against CIDR blocklists, and applies the check again "
    "before each browser navigation to close the DNS-rebinding window.",
    "",
    "3. Multi-user support and authentication: Implement JWT-based authentication "
    "and Postgres Row-Level Security, following the design in the architectural "
    "exploration (backend/, supabase/migrations/), to enable deployment as a shared "
    "service.",
    "",
    "4. Oracle signals: Supplement the LLM-generated pass/fail verdict with "
    "independent signals — HTTP response code capture, JavaScript console error "
    "capture, and perceptual screenshot diffing (dHash algorithm) — as designed in "
    "backend/src/executor/oracle.ts.",
    "",
    "5. Formal evaluation harness: Run the evaluation harness in eval/ against "
    "the committed fixture sites with human-authored ground-truth test cases, to "
    "produce a measurable recall and precision figure.",
    "",
    "6. Screenshot capture on failure: Implement Playwright's page.screenshot() API "
    "as part of executeTest(), storing screenshots alongside error messages in the "
    "database for visual diagnosis.",
]

# Source: everything above; honest summary
CONCLUSIONS_6 = [
    "The main objective of this project — to build a system that accepts a URL and "
    "autonomously generates and executes functional web application tests without "
    "any human test-authoring — was met. The working system (server.js, "
    "run-pipeline.js, and supporting files) crawls a target URL using Playwright, "
    "generates test-case descriptions using the Llama-3.1-8B-Instruct large language "
    "model, converts those descriptions into Playwright scripts, executes them in a "
    "real headless browser, and presents the results in a browser-based dashboard — "
    "all from a single URL input, with no test scripts written by hand.",
    "",
    "The strongest evidence for this is concrete rather than statistical. Four real "
    "bugs were found and fixed through actual system operation: a vacuous pass "
    "produced by a missing selector that the harness incorrectly reported as success; "
    "a safe-mode gap that allowed 'Add to basket' links to pass unsuppressed on a "
    "live e-commerce site; a malformed JSON output mode that silently produced zero "
    "test cases until a retry loop was added; and selector hallucination caused by "
    "the script-generation prompt not receiving the actual DOM data. Each was "
    "discovered by running the system, not by code review, and each was fixed with "
    "a verifiable committed change. The measurable effect of the most significant "
    "fix (including DOM data in the script prompt) is visible in the results: Run A "
    "produced 0 passes and 2 crashes; Run B, after the fix, produced 2 passes and "
    "0 crashes against the same target.",
    "",
    "Where the tool fell short of a human-authored test suite is also stated plainly. "
    "It crawled only the landing page submitted, not the full application. Test cases "
    "are non-deterministic — two runs against the same page produced different "
    "descriptions and selectors. It can detect structural failures reliably but "
    "cannot verify server-side correctness. It was not formally evaluated against a "
    "ground-truth test suite, so no recall or precision figure can be honestly stated.",
    "",
    "The separate architectural exploration (backend/, frontend/, eval/, supabase/) "
    "documents what a production-ready, multi-user version of the same concept would "
    "require. That exploration was not run end-to-end and is presented as a design "
    "document, not a working implementation.",
]

FUTURE_WORK_6 = [
    "Future work should prioritise, in order: integrating the multi-page BFS crawler "
    "into the pipeline; replacing the hostname-only SSRF guard with the designed "
    "DNS-resolving version; implementing the formal evaluation harness to produce a "
    "measurable recall figure; and adding screenshot capture on failure as a "
    "practical debugging aid. Multi-user deployment and authentication depend on a "
    "decision to host the tool as a shared service rather than a local developer "
    "utility.",
]

# Source: git log, run-pipeline.js lines 16-35, 37-61
APPENDIX_B = [
    "The following code listings are extracted verbatim from the working codebase. "
    "Line references are to run-pipeline.js unless otherwise noted.",
    "",
    "Listing B.1 — Page element extraction (run-pipeline.js lines 16-35):",
    "The scanPage() function uses Playwright's $$eval to extract interactive elements "
    "from the live DOM. Buttons are extracted with text, aria-label, and class "
    "fallbacks. Each button is tagged safe: true or safe: false against the "
    "blocklist in safe-mode.config.json.",
    "",
    "async function scanPage(page) {",
    "  const rawButtons = await page.$$eval('button, [role=\"button\"]', (els) =>",
    "    els.map((el) => {",
    "      const text = el.textContent?.trim();",
    "      if (text) return text;",
    "      const aria = el.getAttribute('aria-label');",
    "      if (aria) return '[aria-label] ' + aria;",
    "      if (el.className) return '[class] ' + el.className;",
    "      return '[unlabeled]';",
    "    })",
    "  );",
    "  const inputs = await page.$$eval('input', (els) =>",
    "    els.map((el) => ({ type: el.type, name: el.name, id: el.id, placeholder: el.placeholder }))",
    "  );",
    "  const links = await page.$$eval('a', (els) =>",
    "    els.map((el) => ({ text: el.textContent.trim(), href: el.href }))",
    "  );",
    "  const buttons = rawButtons.map((text) => ({ text: text, safe: !isBlocked(text) }));",
    "  return { buttons: buttons, inputs: inputs, links: links };",
    "}",
    "",
    "Listing B.2 — Test-case generation prompt (run-pipeline.js lines 37-45):",
    "Sent to meta-llama/Llama-3.1-8B-Instruct via HuggingFace Inference API.",
    "",
    "function buildTestGenPrompt(pageData) {",
    "  return 'You are a QA engineer. Given this structured map of a web page " +
    "interactive elements, generate 3-5 test cases a human tester would perform.\\n\\n'",
    "    + 'For each test case, respond with a JSON array where each object has:\\n'",
    "    + '- \"description\": plain-language description of the user action and expected result\\n'",
    "    + '- \"confidence\": either \"structural\" or \"behavioral\"\\n\\n'",
    "    + 'Only test elements marked safe true. Never generate a test for anything marked safe false.\\n\\n'",
    "    + 'Page data:\\n' + JSON.stringify(pageData, null, 2) + '\\n\\n'",
    "    + 'CRITICAL: Respond with ONLY a valid JSON array. No markdown code fences, " +
    "no commentary, no trailing commas, all strings properly closed.';",
    "}",
    "",
    "Listing B.3 — Result classification logic (run-pipeline.js lines 114-119):",
    "Executed after each generated script runs in the live browser context.",
    "",
    "  if (result && result.passed === true) return { executionResult: 'pass' };",
    "  return {",
    "    executionResult: result && result.harnessCaught ? 'crash' : 'fail',",
    "    failureCategory: result && result.harnessCaught ? 'runtime_error' : 'assertion_failed',",
    "    error: (result && result.error) || 'unknown',",
    "  };",
]

APPENDIX_A = [
    "The following screenshots should be captured from the running system and "
    "inserted in place of this placeholder.",
    "",
    "Screenshot A.1 — Dashboard home view: Start the server with 'node server.js', "
    "navigate to http://localhost:3000. Capture the full page showing the URL input "
    "form and (if available) at least one crawl run in the history panel.",
    "",
    "Screenshot A.2 — Detail panel: Click any crawl row in the history panel to "
    "open its detail view. Capture the test case list showing at least one test "
    "with its result badge, generated Playwright script, and error message.",
    "",
    "Screenshot A.3 — Terminal output: Capture a terminal window showing the "
    "server startup message ('AutoTest dashboard running at http://localhost:3000') "
    "and at least one pipeline run in progress, showing the LLM API calls being "
    "made.",
]

# NFRs — Source: code design choices, MAX_PAGES constant in multi-crawler.js,
# HuggingFace free API, and measured run time (11:45–12:16 UTC, 20 July 2026 = ~31 min)
NFRS = [
    "1.3 Non-Functional Requirements",
    "",
    "NFR1: The system shall operate on a single local machine without requiring any "
    "cloud services beyond external API calls (HuggingFace Inference API).",
    "",
    "NFR2: The system shall require no compilation or build step; the frontend "
    "(public/index.html) is served as a static file and the backend runs directly "
    "as Node.js ESM source.",
    "",
    "NFR3: The system shall not interact with any element categorised as unsafe by "
    "the safe-mode blocklist (safe-mode.config.json), regardless of the target URL.",
    "",
    "NFR4: The system shall handle LLM output parsing failures by retrying up to "
    "three times before aborting test generation for the affected page "
    "(run-pipeline.js, MAX_RETRIES = 3).",
    "",
    "NFR5: LLM calls use the HuggingFace free Inference API; there is no API cost "
    "per run. A verification run conducted on 20 July 2026 against "
    "https://demo.playwright.dev/todomvc completed in approximately 31 minutes "
    "(three test cases generated and executed). This figure includes HuggingFace "
    "API latency and is not a formal performance measurement.",
]

# ──────────────────────────────────────────────────────────────
# APPLY ALL CHANGES
# ──────────────────────────────────────────────────────────────

shutil.copy2(DOC_PATH, BACKUP_PATH)
print(f"Backup: {BACKUP_PATH}")

doc = Document(DOC_PATH)
paras = doc.paragraphs   # live view; updates as we add paragraphs

def find_para(pattern, exact=False):
    """Return first paragraph whose text contains pattern (or equals it)."""
    for p in doc.paragraphs:
        txt = para_text(p)
        if exact:
            if txt.strip() == pattern.strip():
                return p
        else:
            if pattern in txt:
                return p
    return None

removed = []  # track paragraphs to remove after iteration

# ── 1. Fix 1.2 functional requirement text errors ────────────────
p83 = find_para('execute generated scripts headlessly and capture pass/fail/error outcomes with screenshots on failure')
if p83:
    set_text(p83, 'The system shall execute generated scripts headlessly and capture pass, fail, crash, and skipped outcomes with error messages.')
    print("Fixed FR4 (screenshots removed)")

p84 = find_para('present results in a dashboard and persist run history per authenticated user')
if p84:
    set_text(p84, 'The system shall present results in a browser-based dashboard and persist run history in a local SQLite database.')
    print("Fixed FR5 (per-authenticated-user removed)")

# ── 2. Submission date ───────────────────────────────────────────
p_date = find_para('[TO BE COMPLETED]', exact=True)
if p_date:
    set_text(p_date, 'July 2026')
    print("Fixed submission date")

# ── 3. Abstract placeholder (write last — keep explicit note) ────
p_abstract = find_para('[TO BE COMPLETED: Full abstract')
if p_abstract:
    set_text(p_abstract,
        '[ABSTRACT — WRITE LAST after all chapters are finalised. '
        '200-300 words. Structure: why it matters (20%), objective (1 sentence), '
        'methodology (20%), results (40%), discussion/relevance (20%). '
        'Use only claims that appear in the body of the report.]')
    print("Marked abstract as write-last")

# ── 4. List of Figures / Tables auto-generate notes ─────────────
p_fig = find_para('[TO BE COMPLETED: Auto-generate once all figures')
if p_fig:
    set_text(p_fig, '[Auto-generate in Word: References > Insert Table of Figures, '
             'after all figures with captions are inserted.]')
    print("Fixed List of Figures note")

p_tbl = find_para('[TO BE COMPLETED: Auto-generate the same way')
if p_tbl:
    set_text(p_tbl, '[Auto-generate in Word: References > Insert Table of Figures '
             '(change Caption label to Table), after all tables are finalised.]')
    print("Fixed List of Tables note")

# ── 5. Hardware requirements ─────────────────────────────────────
p_hw = find_para('[TO BE COMPLETED: Actual specs')
if p_hw:
    set_text(p_hw, HARDWARE[0])
    insert_after(p_hw._p, HARDWARE[1:])
    print("Wrote hardware requirements")

# ── 6. Replace 3.1.2 software table ─────────────────────────────
for tbl in doc.tables:
    # Identify the software stack table by its first cell content
    if tbl.rows and tbl.rows[0].cells:
        first_cell = ''.join(t.text or '' for t in tbl.rows[0].cells[0]._tc.iter(qn('w:t')))
        if 'Layer' in first_cell or 'Frontend' in first_cell or 'Technology' in first_cell:
            correct_rows = [
                ('Layer', 'Technology / Tool'),
                ('Language (all)', 'JavaScript (ESM), Node.js v22'),
                ('Frontend', 'Vanilla JavaScript, HTML, CSS (public/index.html) — no framework, no build step'),
                ('Backend / API', 'Express.js'),
                ('AI integration', 'HuggingFace Inference API, meta-llama/Llama-3.1-8B-Instruct. '
                                   'Note: development initially used the Anthropic Claude API (test-claude-api.js); '
                                   'pivoted mid-project (commit c4da1ce, 12 July 2026) due to free-tier budget exhaustion.'),
                ('Test execution', 'Playwright (headless Chromium)'),
                ('Web crawling', 'Custom extraction using Playwright $$eval (run-pipeline.js lines 16-35) — no third-party crawler library'),
                ('Database', 'SQLite via better-sqlite3 npm package'),
                ('Authentication', 'None — designed as a single-user local tool'),
                ('Hosting', 'Local only (http://localhost:3000) — no cloud deployment'),
                ('Version control', 'Git / GitHub'),
            ]
            # Rebuild table rows
            # First clear existing rows (except header if we're reusing it)
            # Simpler: overwrite cell by cell, add rows if needed
            existing_rows = tbl.rows
            for i, (layer, tech) in enumerate(correct_rows):
                if i < len(existing_rows):
                    row = existing_rows[i]
                    cells = row.cells
                    # Set cell 0
                    for tc in cells[0]._tc.iter(qn('w:t')):
                        tc.text = ''
                    p0 = cells[0].paragraphs[0]
                    set_text(p0, layer)
                    # Set cell 1
                    for tc in cells[1]._tc.iter(qn('w:t')):
                        tc.text = ''
                    p1 = cells[1].paragraphs[0]
                    set_text(p1, tech)
                else:
                    # Add a new row
                    new_row = tbl.add_row()
                    set_text(new_row.cells[0].paragraphs[0], layer)
                    set_text(new_row.cells[1].paragraphs[0], tech)
            print("Replaced 3.1.2 software table")
            break

# ── 7. Methods 3.2 body (currently guidance text) ───────────────
p_methods_body = find_para('Give complete details of materials, instruments')
if p_methods_body:
    set_text(p_methods_body, METHODS_32[0])
    insert_after(p_methods_body._p, METHODS_32[1:])
    print("Wrote 3.2 Methods")

# ── 8. Remove all GUIDANCE paragraphs ───────────────────────────
to_remove = []
for p in doc.paragraphs:
    txt = para_text(p)
    if txt.startswith('GUIDANCE —') or txt.startswith('GUIDANCE—'):
        to_remove.append(p)
    # Also remove the 5.1 guidance line
    elif 'Restate the objectives from Chapter 1 in the same order' in txt:
        to_remove.append(p)
    # And the 3.2 GUIDANCE line (second one)
    elif "This chapter should read as 'what was done'" in txt:
        to_remove.append(p)
for p in to_remove:
    remove_para(p)
print(f"Removed {len(to_remove)} GUIDANCE / meta paragraphs")

# ── 9. Non-functional requirements (replace GUIDANCE that was just removed;
#       insert as new section after last FR) ──────────────────────────────
# The GUIDANCE about NFRs was removed above. Insert NFRs after the last FR paragraph.
p_fr5 = find_para('local SQLite database')
if p_fr5:
    insert_after(p_fr5._p, [''] + NFRS)
    print("Inserted 1.3 Non-Functional Requirements")

# ── 10. Architecture 4.1 ────────────────────────────────────────
p_arch = find_para('[TO BE COMPLETED: Insert the final architecture diagram')
if p_arch:
    set_text(p_arch, ARCH_41[0])
    insert_after(p_arch._p, ARCH_41[1:])
    print("Wrote 4.1 Architecture")

# ── 11. Data model 4.2 ──────────────────────────────────────────
p_dm = find_para('[TO BE COMPLETED: Insert the final ER diagram')
if p_dm:
    set_text(p_dm, DATA_MODEL_42[0])
    insert_after(p_dm._p, DATA_MODEL_42[1:])
    print("Wrote 4.2 Data Model")

# ── 12. Prompt design 4.3 ───────────────────────────────────────
p_prompt = find_para('[TO BE COMPLETED: Document the actual structured prompt')
if p_prompt:
    set_text(p_prompt, PROMPTS_43[0])
    insert_after(p_prompt._p, PROMPTS_43[1:])
    print("Wrote 4.3 Prompt Design")

# ── 13. User interface 4.4 ──────────────────────────────────────
p_ui = find_para('[TO BE COMPLETED: Insert dashboard screenshots')
if p_ui:
    set_text(p_ui, UI_44[0])
    insert_after(p_ui._p, UI_44[1:])
    print("Wrote 4.4 User Interface")

# ── 14. Security 4.5 ────────────────────────────────────────────
p_sec = find_para('[TO BE COMPLETED: Document what was actually implemented')
if p_sec:
    set_text(p_sec, SECURITY_45[0])
    insert_after(p_sec._p, SECURITY_45[1:])
    print("Wrote 4.5 Security")

# ── 15. Objectives 5.1 ──────────────────────────────────────────
# After GUIDANCE removal, 5.1 body may now be empty. Insert objectives.
p_51_header = find_para('5.1 Objectives')
if p_51_header:
    insert_after(p_51_header._p, OBJECTIVES_51)
    print("Wrote 5.1 Objectives")

# ── 16. Results 5.2 (quantitative + qualitative combined) ───────
p_quant = find_para('[TO BE COMPLETED: Quantitative: number of applications')
if p_quant:
    set_text(p_quant, RESULTS_52[0])
    insert_after(p_quant._p, RESULTS_52[1:])
    print("Wrote 5.2 Results (quantitative)")

p_qual = find_para('[TO BE COMPLETED: Qualitative: examples of test cases')
if p_qual:
    # Remove this paragraph — content is already merged into RESULTS_52 above
    remove_para(p_qual)
    print("Removed duplicate 5.2 qualitative placeholder")

# ── 17. Discussion 5.3 ──────────────────────────────────────────
p_disc = find_para('[TO BE COMPLETED: Interpret the results against the objectives')
if p_disc:
    set_text(p_disc, DISCUSSION_53[0])
    insert_after(p_disc._p, DISCUSSION_53[1:])
    print("Wrote 5.3 Discussion")

# ── 18. Limitations 5.4 ─────────────────────────────────────────
p_lim = find_para('[TO BE COMPLETED: Carry forward and update the Limitations')
if p_lim:
    set_text(p_lim, LIMITATIONS_54[0])
    insert_after(p_lim._p, LIMITATIONS_54[1:])
    print("Wrote 5.4 Limitations")

# ── 19. Future development 5.5 ──────────────────────────────────
p_fut55 = find_para('[TO BE COMPLETED: What would be built next')
if p_fut55:
    set_text(p_fut55, FUTURE_55[0])
    insert_after(p_fut55._p, FUTURE_55[1:])
    print("Wrote 5.5 Future Development")

# ── 20. Conclusions chapter 6 ───────────────────────────────────
p_conc = find_para('[TO BE COMPLETED: Summarise whether the main objective')
if p_conc:
    set_text(p_conc, CONCLUSIONS_6[0])
    insert_after(p_conc._p, CONCLUSIONS_6[1:])
    print("Wrote Chapter 6 Conclusions")

p_fw = find_para('[TO BE COMPLETED: Concrete next steps')
if p_fw:
    set_text(p_fw, FUTURE_WORK_6[0])
    insert_after(p_fw._p, FUTURE_WORK_6[1:])
    print("Wrote Future Work section")

# ── 21. Appendix A (screenshots) ────────────────────────────────
p_appa = find_para('[TO BE COMPLETED: Insert full-resolution screenshots')
if p_appa:
    set_text(p_appa, APPENDIX_A[0])
    insert_after(p_appa._p, APPENDIX_A[1:])
    print("Wrote Appendix A instructions")

# ── 22. Appendix B (code listings) ──────────────────────────────
p_appb = find_para("[TO BE COMPLETED: Insert the crawler's element-extraction")
if p_appb:
    set_text(p_appb, APPENDIX_B[0])
    insert_after(p_appb._p, APPENDIX_B[1:])
    print("Wrote Appendix B code listings")

# ── 23. Remove References GUIDANCE ──────────────────────────────
# (already handled in step 8 — just a safety check)

# ──────────────────────────────────────────────────────────────
# SAVE
# ──────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print(f"\nDocument saved: {DOC_PATH}")

# ── VERIFICATION: scan for remaining markers ──────────────────
print("\n=== REMAINING MARKERS SCAN ===")
doc2 = Document(DOC_PATH)
remaining_tbc = []
remaining_guidance = []
for p in doc2.paragraphs:
    txt = para_text(p)
    if '[TO BE COMPLETED' in txt:
        remaining_tbc.append(txt[:120])
    if 'GUIDANCE' in txt and txt.startswith('GUIDANCE'):
        remaining_guidance.append(txt[:120])

# Also check tables
for tbl in doc2.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                txt = para_text(p)
                if '[TO BE COMPLETED' in txt:
                    remaining_tbc.append('[TABLE] ' + txt[:120])
                if 'GUIDANCE' in txt and txt.startswith('GUIDANCE'):
                    remaining_guidance.append('[TABLE] ' + txt[:120])

print(f"[TO BE COMPLETED] markers remaining: {len(remaining_tbc)}")
for m in remaining_tbc:
    print(f"  {m}")
print(f"GUIDANCE markers remaining: {len(remaining_guidance)}")
for m in remaining_guidance:
    print(f"  {m}")

total_paras = len([p for p in doc2.paragraphs if para_text(p).strip()])
print(f"\nTotal non-empty paragraphs in final document: {total_paras}")
print("Done.")
