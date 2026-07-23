#!/usr/bin/env python3
"""
AutoTest Final Report — comprehensive update script.
Handles: Steps 2-6 from the audit.
  - Insert 6 images (Figures 1-4, Appendix A screenshots)
  - Fix abstract
  - Add Codebase B precision sentence
  - Fix heading styles
  - Insert real TOC / List of Figures / List of Tables field codes
  - Monospace font for Appendix B code listings
  - Fix test count claim (30 -> 57)
  - Report on Declaration blanks
"""

import shutil, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from copy import deepcopy

DOC_PATH    = '/home/madushan/Project/autotest/AutoTest_Final_Report.docx'
BACKUP_PATH = '/home/madushan/Project/autotest/AutoTest_Final_Report_Step2_BACKUP.docx'
ASSETS      = '/home/madushan/Project/autotest/report_assets'

shutil.copy2(DOC_PATH, BACKUP_PATH)
print(f"Backup: {BACKUP_PATH}")

doc = Document(DOC_PATH)

# ──────────────────────────────────────────────────────────────────
# UTILITIES
# ──────────────────────────────────────────────────────────────────

def para_text(para):
    return ''.join(t.text or '' for t in para._p.iter(qn('w:t')))

def find_para(needle):
    for p in doc.paragraphs:
        if needle in para_text(p):
            return p
    return None

def set_text(para, text):
    p = para._p
    for r in list(p.iter(qn('w:r'))):
        r.getparent().remove(r)
    for hl in list(p.iter(qn('w:hyperlink'))):
        hl.getparent().remove(hl)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)

def make_p(text='', style_id=None):
    new_p = OxmlElement('w:p')
    if style_id:
        pPr = OxmlElement('w:pPr')
        ps = OxmlElement('w:pStyle')
        ps.set(qn('w:val'), style_id)
        pPr.append(ps)
        new_p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        new_p.append(r)
    return new_p

def insert_after(ref_elem, lines, style_id=None):
    last = ref_elem
    for text in lines:
        new_p = make_p(text, style_id)
        last.addnext(new_p)
        last = new_p

def remove_para(para):
    p = para._p
    p.getparent().remove(p)

def add_run_to_para(para, text, font_name=None, font_size=None, bold=False,
                    italic=False, color=None):
    run = para.add_run(text)
    if font_name:  run.font.name = font_name
    if font_size:  run.font.size = Pt(font_size)
    if bold:       run.font.bold = bold
    if italic:     run.font.italic = italic
    if color:      run.font.color.rgb = RGBColor(*color)
    return run

# ──────────────────────────────────────────────────────────────────
# STEP 2A: Apply proper Heading styles to chapter/section headings
# ──────────────────────────────────────────────────────────────────
# Map paragraph text prefix -> heading level
HEADING_PATTERNS = [
    # (pattern, heading_level)
    (r'^Chapter \d+:', 1),
    (r'^\d+\.\d+\.\d+ ', 3),
    (r'^\d+\.\d+ ', 2),
    (r'^Appendix [A-Z]:', 1),
    (r'^References$', 1),
    (r'^Declaration$', 1),
    (r'^Abstract$', 1),
    (r'^Acknowledgements$', 1),
    (r'^Table of Contents$', 1),
    (r'^List of Figures$', 1),
    (r'^List of Tables$', 1),
    (r'^List of Abbreviations$', 1),
    (r'^Future Work and Recommendations$', 1),
]

heading_style_map = {1: 'Heading1', 2: 'Heading2', 3: 'Heading3'}

fixed_headings = 0
for para in doc.paragraphs:
    txt = para_text(para).strip()
    for pattern, level in HEADING_PATTERNS:
        if re.match(pattern, txt):
            style_name = f'Heading {level}'
            # Check if style exists in doc
            try:
                para.style = doc.styles[style_name]
                fixed_headings += 1
            except KeyError:
                # Style doesn't exist; set via XML
                pPr = para._p.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    para._p.insert(0, pPr)
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is None:
                    pStyle = OxmlElement('w:pStyle')
                    pPr.insert(0, pStyle)
                pStyle.set(qn('w:val'), f'Heading{level}')
                fixed_headings += 1
            break

print(f"Fixed {fixed_headings} heading styles")

# ──────────────────────────────────────────────────────────────────
# STEP 2B: Insert images replacing [Figure N — INSERT HERE] placeholders
# ──────────────────────────────────────────────────────────────────

def insert_image_replacing_para(placeholder_text, image_path, caption_text, width_inches=5.5):
    """
    Find paragraph containing placeholder_text.
    Replace it with: [image paragraph] + [caption paragraph using Caption style].
    Returns True if found and replaced.
    """
    para = find_para(placeholder_text)
    if not para:
        print(f"  WARNING: placeholder not found: {placeholder_text[:60]}")
        return False

    ref_elem = para._p

    # 1. Create caption paragraph AFTER (insert first since addnext reverses)
    cap_p = make_p()
    # Try to apply Caption style
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Caption')
    pPr.append(pStyle)
    cap_p.append(pPr)
    cap_r = OxmlElement('w:r')
    cap_t = OxmlElement('w:t')
    cap_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    cap_t.text = caption_text
    cap_r.append(cap_t)
    cap_p.append(cap_r)
    ref_elem.addnext(cap_p)

    # 2. Create image paragraph BEFORE the caption (i.e., addnext on ref_elem again)
    img_p = OxmlElement('w:p')
    # Center the paragraph
    pPr2 = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr2.append(jc)
    img_p.append(pPr2)

    img_r = OxmlElement('w:r')
    drawing = doc.inline_shapes._body  # we'll use add_picture differently

    # Use python-docx paragraph API: add to a temp paragraph then move XML
    temp_para = doc.add_paragraph()
    run = temp_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    # Get the drawing XML from the temp para's run
    drawing_xml = temp_para._p.find(f'.//{qn("w:drawing")}')
    if drawing_xml is None:
        # Try alternate namespace
        drawing_xml = temp_para._p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')

    if drawing_xml is not None:
        img_r.append(deepcopy(drawing_xml))

    img_p.append(img_r)

    # Remove the temp paragraph
    temp_para._p.getparent().remove(temp_para._p)

    # Insert img_p after ref_elem (before cap_p since cap_p was inserted first)
    ref_elem.addnext(img_p)

    # 3. Clear the original placeholder paragraph (blank it)
    set_text(para, '')

    return True


IMAGE_SPECS = [
    # (placeholder_text, image_file, caption, width_inches)
    ('[Figure 1 — INSERT HERE',
     f'{ASSETS}/figure1_architecture.png',
     'Figure 1: AutoTest system architecture (run-pipeline.js, server.js). '
     'Dashed-border boxes denote external systems.',
     5.5),
    ('[Figure 2 — INSERT HERE',
     f'{ASSETS}/figure2_er_diagram.png',
     'Figure 2: SQLite database schema (db.js lines 5–32). '
     'All timestamps stored as ISO 8601 TEXT strings.',
     5.5),
    ('[Figure 3 — INSERT HERE',
     f'{ASSETS}/figure3_dashboard_home.png',
     'Figure 3: AutoTest dashboard — URL input form and crawl history panel '
     'showing one run for https://demo.playwright.dev/todomvc (2 pass, 1 fail, 0 crash).',
     5.5),
    ('[Figure 4 — INSERT HERE',
     f'{ASSETS}/figure4_detail_panel.png',
     'Figure 4: Detail view for crawl #1 — three generated test cases with results. '
     'Two tests passed; one failed due to a strict-mode locator ambiguity.',
     5.5),
]

for placeholder, img_path, caption, width in IMAGE_SPECS:
    ok = insert_image_replacing_para(placeholder, img_path, caption, width)
    print(f"  {'✓' if ok else '✗'} Image inserted: {caption[:60]}")

# ── Appendix A screenshots ─────────────────────────────────────
appendix_specs = [
    ('[Screenshot A.1',
     f'{ASSETS}/figure3_dashboard_home.png',
     'Screenshot A.1: Dashboard home view — URL input form and crawl history.',
     5.2),
    ('[Screenshot A.2',
     f'{ASSETS}/figure4_detail_panel.png',
     'Screenshot A.2: Detail panel — test case results with FAIL/PASS badges and error messages.',
     5.2),
    ('[Screenshot A.3',
     f'{ASSETS}/appendix_a1_full_view.png',
     'Screenshot A.3: Full dashboard view at 1440px — history and detail panel visible simultaneously.',
     5.2),
]

for placeholder, img_path, caption, width in appendix_specs:
    ok = insert_image_replacing_para(placeholder, img_path, caption, width)
    print(f"  {'✓' if ok else '✗'} Appendix screenshot: {caption[:60]}")

# ──────────────────────────────────────────────────────────────────
# STEP 3: Fix the abstract
# ──────────────────────────────────────────────────────────────────
ABSTRACT_TEXT = (
    "Manual test authoring represents a significant time and skill barrier for solo developers, "
    "freelancers, and student teams without dedicated quality assurance resources. This report "
    "documents the design, implementation, and evaluation of AutoTest, a zero-configuration "
    "testing agent that removes the requirement to author web application test scripts by hand.\n\n"
    "The objective was to build a system that accepts a single URL and autonomously generates and "
    "executes functional tests against the target application.\n\n"
    "The system was implemented in Node.js using Playwright for browser automation, the HuggingFace "
    "Inference API (meta-llama/Llama-3.1-8B-Instruct) for test generation, SQLite for result "
    "persistence, and a vanilla JavaScript dashboard for result inspection. Development followed an "
    "incremental approach over two days (12-13 July 2026), with each pipeline stage committed to "
    "version control after being verified against a real web application. The initial LLM integration "
    "used the Anthropic Claude API; this was replaced mid-project with the HuggingFace free Inference "
    "API following free-tier credit exhaustion.\n\n"
    "Two pipeline runs were conducted against the same public target "
    "(https://demo.playwright.dev/todomvc). The first run (12 July 2026) produced three test cases "
    "and recorded zero passes, one assertion failure, and two crashes, all attributable to "
    "LLM-generated selector errors. The second run (20 July 2026), conducted after a prompt fix "
    "requiring the script-generation stage to use actual crawled DOM data, produced two passes and "
    "one assertion failure, with zero crashes. Four specific bugs were identified and corrected "
    "through real system operation during development, each traceable to a committed code change.\n\n"
    "The results demonstrate that the execution engine reliably distinguishes runtime errors from "
    "assertion failures in AI-generated test code. A separate finding - a test that returned a false "
    "positive by iterating zero times over a hallucinated selector - was identified through manual "
    "result review rather than automated detection, and is discussed as an open limitation rather "
    "than a solved capability. The principal remaining limitations are LLM output non-determinism, "
    "a single-page crawl scope in the integrated pipeline, and the absence of a formal evaluation "
    "against a ground-truth test suite. The tool operates as a local single-user prototype; a "
    "production deployment would require DNS-resolving SSRF protection, authentication, and a formal "
    "evaluation harness, all of which were designed in a separate architectural exploration but "
    "not integrated."
)

p_abstract = find_para('[ABSTRACT — WRITE LAST')
if p_abstract:
    set_text(p_abstract, ABSTRACT_TEXT)
    print("✓ Step 3: Abstract written")
else:
    print("✗ Step 3: Abstract placeholder not found")

# ──────────────────────────────────────────────────────────────────
# STEP 4: Codebase B precision fix
# ──────────────────────────────────────────────────────────────────
PRECISION_SENTENCE = (
    "Note on Codebase B integration status: individual components of the separate architectural "
    "exploration were run and debugged in isolation during its development — for example, the "
    "login-redirect flow and CORS/authentication handling were exercised (commits of 15 July 2026). "
    "What was never exercised is the full integrated pipeline end-to-end: crawl through Supabase "
    "persistence through the React dashboard as a complete connected flow. Both statements are "
    "true and not contradictory: component-level debugging occurred; full-pipeline integration did not."
)

# Find the 'never run end-to-end' sentence and add after it
targets = [
    'never run end-to-end',
    'not run end-to-end',
    'presented as a design document',
]
precision_added = False
for t in targets:
    p = find_para(t)
    if p:
        insert_after(p._p, ['', PRECISION_SENTENCE])
        print(f"✓ Step 4: Codebase B precision sentence added after para containing '{t[:40]}'")
        precision_added = True
        break
if not precision_added:
    print("✗ Step 4: Could not find 'never run end-to-end' paragraph")

# ──────────────────────────────────────────────────────────────────
# STEP 5: TOC / List of Figures / List of Tables field codes
# (Word field codes — will auto-update when opened in Word)
# ──────────────────────────────────────────────────────────────────

def make_field_para(field_code, instruction_text=None):
    """Create a paragraph containing a Word field code (e.g. TOC)."""
    new_p = OxmlElement('w:p')
    # Simple field using w:fldSimple
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), field_code)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '[Right-click and select Update Field to generate in Word]'
    r.append(t)
    fld.append(r)
    new_p.append(fld)
    return new_p

# Replace TOC placeholder
toc_placeholder = find_para('[Auto-generate in Word: References > Insert Table of Figures,\nafter all figures with captions are inserted.]')
if not toc_placeholder:
    toc_placeholder = find_para('[Auto-generate in Word: References > Insert Table of Figures')

# Actually find the table of contents section
toc_section = find_para('Table of Contents')
if toc_section:
    # Insert a TOC field after the heading
    toc_field = make_field_para('TOC \\o "1-3" \\h \\z \\u')
    toc_section._p.addnext(toc_field)
    print("✓ Step 5: TOC field inserted after 'Table of Contents' heading")

# Replace List of Figures placeholder
lof_placeholder = find_para('[Auto-generate in Word: References > Insert Table of Figures,')
if lof_placeholder:
    lof_field = make_field_para('TOC \\h \\z \\c "Figure"')
    lof_placeholder._p.addnext(lof_field)
    set_text(lof_placeholder, '')
    print("✓ Step 5: List of Figures field inserted")

# Replace List of Tables placeholder
lot_placeholder = find_para('[Auto-generate in Word: References > Insert Table of Figures\n(change Caption label to Table)')
if not lot_placeholder:
    lot_placeholder = find_para('(change Caption label to Table)')
if lot_placeholder:
    lot_field = make_field_para('TOC \\h \\z \\c "Table"')
    lot_placeholder._p.addnext(lot_field)
    set_text(lot_placeholder, '')
    print("✓ Step 5: List of Tables field inserted")

# ──────────────────────────────────────────────────────────────────
# STEP 6A: Monospace font for Appendix B code listings
# ──────────────────────────────────────────────────────────────────

CODE_MARKERS = [
    'async function scanPage(',
    'function buildTestGenPrompt(',
    'function buildScriptGenPrompt(',
    'const rawButtons',
    'const inputs = await',
    'const links = await',
    'const buttons = rawButtons',
    'return { buttons',
    'if (result && result.passed',
    'executionResult:',
    'failureCategory:',
    "return 'You are a QA engineer.",
    "+ 'For each test case",
    '+ JSON.stringify(',
    "return 'You are a Playwright",
    '+ JSON.stringify(pageData',
    '  if (result && result.passed',
]

mono_fixed = 0
for para in doc.paragraphs:
    txt = para_text(para)
    is_code = any(marker in txt for marker in CODE_MARKERS)
    # Also detect lines in Appendix B that look like code (start with spaces or specific patterns)
    if not is_code:
        stripped = txt.strip()
        is_code = (
            stripped.startswith('await ') or
            stripped.startswith('const ') or
            stripped.startswith('return ') or
            stripped.startswith('};') or
            stripped.startswith('}') or
            stripped.startswith("  if") or
            stripped.startswith("  return") or
            (stripped.startswith('{') and 'passed' in stripped) or
            stripped.startswith("+ '") or
            stripped.startswith('+ JSON') or
            (stripped.startswith('"') and 'passed' in stripped)
        )

    if is_code:
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        # Apply to all runs including those we can't access via .runs
        for r_elem in para._p.iter(qn('w:r')):
            rPr = r_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r_elem.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), 'Courier New')
            rFonts.set(qn('w:hAnsi'), 'Courier New')
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                rPr.append(sz)
            sz.set(qn('w:val'), '18')  # 9pt = 18 half-points
        mono_fixed += 1

print(f"✓ Step 6A: Applied Courier New to {mono_fixed} code-block paragraphs in Appendix B")

# ──────────────────────────────────────────────────────────────────
# STEP 6B: Fix test count claim (30 -> 57 with breakdown)
# ──────────────────────────────────────────────────────────────────
p_30tests = find_para('30 tests') 
if p_30tests:
    old_text = para_text(p_30tests)
    new_text = old_text.replace(
        '(30 tests)',
        '(57 tests: 30 SSRF-guard tests, 14 retry-logic tests, and 13 safety-filter tests — all passing)'
    ).replace(
        '30 tests',
        '57 tests (30 SSRF-guard tests, 14 retry-logic tests, and 13 safety-filter tests — all passing)'
    )
    set_text(p_30tests, new_text)
    print(f"✓ Step 6B: Test count updated to 57")
else:
    print("  (30-test paragraph not found — may have already been updated)")

# ──────────────────────────────────────────────────────────────────
# STEP 6C: Declaration blank fields — report only
# ──────────────────────────────────────────────────────────────────
print("\n── Step 6C: Declaration / Certification fields ──")
print("  Supervisor Name, Signature, Date lines in the Declaration section")
print("  contain only '___________________________' blank lines.")
print("  These are intentionally blank — they require wet (physical) signatures.")
print("  No name or date has been invented or inserted.")
print("  Action required: print the document and obtain physical signature from supervisor.")

# ──────────────────────────────────────────────────────────────────
# SAVE
# ──────────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print(f"\n✓ Document saved: {DOC_PATH}")

# ──────────────────────────────────────────────────────────────────
# STEP 7 PREP: Count images and scan for remaining placeholders
# ──────────────────────────────────────────────────────────────────
print("\n=== VERIFICATION ===")
doc2 = Document(DOC_PATH)

# Count images in docx XML
import zipfile, os
with zipfile.ZipFile(DOC_PATH, 'r') as z:
    media_files = [f for f in z.namelist() if f.startswith('word/media/')]
    image_files = [f for f in media_files if any(f.endswith(e) for e in ['.png','.jpg','.jpeg','.gif','.bmp'])]
print(f"Images embedded in docx: {len(image_files)}")
for img in image_files:
    print(f"  {img}")

# Scan for remaining placeholders
remaining = []
for p in doc2.paragraphs:
    txt = ''.join(t.text or '' for t in p._p.iter(qn('w:t')))
    if any(marker in txt for marker in ['[TO BE COMPLETED', '[Figure N', '[INSERT HERE', '[INSERT SCREENSHOT', 'GUIDANCE —', '[Auto-generate']):
        remaining.append(txt[:120])
# Also check tables
for tbl in doc2.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                txt = ''.join(t.text or '' for t in p._p.iter(qn('w:t')))
                if any(marker in txt for marker in ['[TO BE COMPLETED', '[Figure N', 'GUIDANCE —', '[Auto-generate']):
                    remaining.append('[TABLE] ' + txt[:120])

print(f"\nRemaining placeholder markers: {len(remaining)}")
for r in remaining:
    print(f"  {r[:100]}")

# Word count
total_words = 0
for p in doc2.paragraphs:
    total_words += len(para_text(p).split())
print(f"\nApproximate word count (body paragraphs): {total_words}")
print(f"Total paragraphs (non-empty): {len([p for p in doc2.paragraphs if para_text(p).strip()])}")

print("\nDone.")
